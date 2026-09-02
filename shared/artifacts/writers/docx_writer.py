from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..exceptions import ArtifactPatchError, ArtifactWriteError
from ..models import ArtifactDocument, ArtifactType, BlockType, ContentBlock
from ..patch import LocationTracker, require_table, require_text
from ..spec import ArtifactSpec, PatchOp, SpecBlock, list_items


class DocxWriter:
    kind = ArtifactType.WORD

    def write(self, spec: ArtifactSpec, dest: Path) -> None:
        try:
            document = Document()
            _set_core(document, spec)
            blocks = list(spec.blocks)
            if not blocks and spec.title:
                blocks = [SpecBlock(type=BlockType.HEADING, text=spec.title, level=1)]
            for block in blocks:
                _append_block(document, block)
            dest.parent.mkdir(parents=True, exist_ok=True)
            document.save(str(dest))
        except ArtifactWriteError:
            raise
        except Exception as exc:
            raise ArtifactWriteError(f"Failed to write Word file {dest.name}: {exc}") from exc

    def apply_ops(
        self,
        path: Path,
        document: ArtifactDocument,
        ops: list[PatchOp],
        dest: Path,
    ) -> None:
        native = Document(str(path))
        tracker = LocationTracker(document.blocks)
        for op in ops:
            _apply_docx_op(native, tracker, op)
        dest.parent.mkdir(parents=True, exist_ok=True)
        native.save(str(dest))


def _set_core(document: Document, spec: ArtifactSpec) -> None:
    if spec.title:
        document.core_properties.title = spec.title
    if spec.author:
        document.core_properties.author = spec.author


def _append_block(document: Document, block: SpecBlock) -> None:
    if block.type == BlockType.HEADING:
        level = min(max(block.level or 1, 1), 9)
        text = block.text.strip()
        if text:
            document.add_heading(text, level=level)
        return
    if block.type == BlockType.LIST:
        items = list_items(block)
        if not items and block.text.strip():
            items = [block.text.strip()]
        for item in items:
            document.add_paragraph(item, style="List Bullet")
        return
    if block.type == BlockType.TABLE:
        rows = block.table or []
        if rows:
            _add_table(document, rows)
        return
    text = block.text.strip()
    if text:
        document.add_paragraph(text)


def _add_table(document: Document, rows: list[list[str]]) -> Table:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    _fill_table(table, rows)
    return table


def _fill_table(table: Table, rows: list[list[str]]) -> None:
    for r_i, row in enumerate(rows):
        for c_i, value in enumerate(row):
            if c_i < len(table.rows[r_i].cells):
                table.rows[r_i].cells[c_i].text = value


def _apply_docx_op(native: Document, tracker: LocationTracker, op: PatchOp) -> None:
    if op.op == "set_title":
        title = op.title or op.text
        if title is None:
            raise ArtifactPatchError("set_title requires 'title' or 'text'.")
        native.core_properties.title = title
        return
    if op.op == "set_cell":
        raise ArtifactPatchError("set_cell is only valid for excel artifacts.")
    if op.op == "insert_slide":
        raise ArtifactPatchError("insert_slide is only valid for ppt artifacts.")
    if op.op == "replace_shape_text":
        raise ArtifactPatchError("replace_shape_text is only valid for ppt artifacts.")
    if op.op == "replace_text":
        block = tracker.find(op)
        idx = block.location.paragraph_index
        if idx is None:
            raise ArtifactPatchError(f"Block {block.block_id} has no paragraph_index to patch.")
        native.paragraphs[idx].text = require_text(op)
        return
    if op.op == "replace_table":
        block = tracker.find(op)
        idx = block.location.table_index
        if idx is None:
            raise ArtifactPatchError(f"Block {block.block_id} has no table_index to patch.")
        _replace_table(native, idx, require_table(op))
        return
    if op.op == "insert_block":
        _insert_block(native, tracker, op)
        return
    if op.op == "delete_block":
        _delete_block(native, tracker, op)
        return
    raise ArtifactPatchError(f"Unknown patch op {op.op!r}.")


def _insert_block(native: Document, tracker: LocationTracker, op: PatchOp) -> None:
    if op.block is None:
        raise ArtifactPatchError("insert_block requires a 'block' payload.")
    new_block = op.block
    if not op.after_block_id and not op.block_id and op.location is None:
        _append_block(native, new_block)
        return
    anchor = tracker.find_anchor(op)
    if new_block.type == BlockType.TABLE and new_block.table:
        _insert_table_after(native, anchor, new_block.table)
        tracker.shift_tables(anchor.location.table_index or -1, 1)
        return
    texts, style = _paragraph_payload(new_block)
    if anchor.type == BlockType.TABLE and anchor.location.table_index is not None:
        ref = native.tables[anchor.location.table_index]._tbl
        last = None
        for text in texts:
            last = _insert_paragraph_xml_after(ref, text, style, native)
            ref = last._p
        tracker.shift_paragraphs(_max_paragraph_index(anchor, native), len(texts))
        return
    idx = anchor.location.paragraph_index
    if idx is None:
        raise ArtifactPatchError("Cannot insert after a block without paragraph_index.")
    paragraph = native.paragraphs[idx]
    for text in texts:
        paragraph = _insert_paragraph_after(paragraph, text, style)
    tracker.shift_paragraphs(idx, len(texts))


def _delete_block(native: Document, tracker: LocationTracker, op: PatchOp) -> None:
    block = tracker.find(op)
    if block.type == BlockType.TABLE and block.location.table_index is not None:
        idx = block.location.table_index
        tbl = native.tables[idx]._tbl
        tbl.getparent().remove(tbl)
        tracker.shift_tables(idx, -1)
        tracker.remove(block)
        return
    idx = block.location.paragraph_index
    if idx is None:
        raise ArtifactPatchError(f"Block {block.block_id} cannot be deleted without paragraph_index.")
    element = native.paragraphs[idx]._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    tracker.shift_paragraphs(idx, -1)
    tracker.remove(block)


def _replace_table(document: Document, table_index: int, rows: list[list[str]]) -> None:
    table = document.tables[table_index]
    n_rows = len(rows)
    n_cols = max((len(row) for row in rows), default=0)
    if n_rows == 0 or n_cols == 0:
        raise ArtifactPatchError("replace_table requires a non-empty 2D table.")
    if len(table.rows) == n_rows and len(table.columns) == n_cols:
        _fill_table(table, rows)
        return
    old = table._tbl
    parent = old.getparent()
    xml_index = parent.index(old)
    parent.remove(old)
    new_table = document.add_table(rows=n_rows, cols=n_cols)
    try:
        new_table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    new_tbl = new_table._tbl
    new_tbl.getparent().remove(new_tbl)
    parent.insert(xml_index, new_tbl)
    _fill_table(new_table, rows)


def _insert_table_after(document: Document, anchor: ContentBlock, rows: list[list[str]]) -> None:
    table = _add_table(document, rows)
    new_tbl = table._tbl
    new_tbl.getparent().remove(new_tbl)
    if anchor.type == BlockType.TABLE and anchor.location.table_index is not None:
        ref = document.tables[anchor.location.table_index]._tbl
    elif anchor.location.paragraph_index is not None:
        ref = document.paragraphs[anchor.location.paragraph_index]._p
    else:
        raise ArtifactPatchError("Cannot insert table after block without a native anchor.")
    ref.addnext(new_tbl)


def _paragraph_payload(block: SpecBlock) -> tuple[list[str], str | None]:
    if block.type == BlockType.HEADING:
        level = min(max(block.level or 1, 1), 9)
        return [block.text.strip()], f"Heading {level}"
    if block.type == BlockType.LIST:
        return list_items(block) or [block.text.strip()], "List Bullet"
    return [block.text.strip()], None


def _insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def _insert_paragraph_xml_after(ref, text: str, style: str | None, document: Document) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref.addnext(new_p)
    new_para = Paragraph(new_p, document)
    if style:
        new_para.style = style
    new_para.text = text
    return new_para


def _max_paragraph_index(anchor: ContentBlock, native: Document) -> int:
    if anchor.location.paragraph_index is not None:
        return anchor.location.paragraph_index
    return max((i for i, _p in enumerate(native.paragraphs)), default=0)
