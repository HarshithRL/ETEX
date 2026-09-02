from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from ..blocks import make_content_block
from ..exceptions import ArtifactParseError, CorruptArtifact
from ..heading_stack import HeadingStack
from ..ids import file_artifact_id
from ..models import (
    ArtifactDocument,
    ArtifactMetadata,
    ArtifactType,
    BlockType,
    ContentBlock,
    MIME_BY_TYPE,
    ParseOptions,
    SourceLocation,
)
from .pdf_tables import table_to_text

_HEADING_STYLE = re.compile(r"heading\s*(\d+)", re.I)


class DocxParser:
    kind = ArtifactType.WORD

    def parse(self, path: Path, options: ParseOptions) -> ArtifactDocument:
        try:
            document = Document(str(path))
        except PackageNotFoundError as exc:
            raise CorruptArtifact(f"Unreadable Word file: {path.name}") from exc
        except Exception as exc:
            raise CorruptArtifact(f"Failed to open Word file: {path.name}") from exc

        core = document.core_properties
        blocks: list[ContentBlock] = []
        warnings: list[str] = []
        index = 0
        stack = HeadingStack()

        try:
            for para_i, paragraph in enumerate(document.paragraphs):
                text = (paragraph.text or "").strip()
                if not text:
                    continue
                style_name = (paragraph.style.name if paragraph.style else "") or ""
                block_type = (
                    BlockType.HEADING
                    if style_name.lower().startswith("heading")
                    else BlockType.TEXT
                )
                if style_name.lower().startswith("list"):
                    block_type = BlockType.LIST
                level = None
                heading_path: list[str] = stack.path()
                if block_type == BlockType.HEADING:
                    match = _HEADING_STYLE.search(style_name)
                    level = int(match.group(1)) if match else 1
                    heading_path = stack.push(level, text.split("\n", 1)[0])
                index += 1
                extra = {"style": style_name} if style_name else {}
                blocks.append(
                    make_content_block(
                        seq_id=f"word-{index:04d}",
                        block_type=block_type,
                        text=text,
                        location=SourceLocation(paragraph_index=para_i),
                        extra=extra,
                        level=level,
                        heading_path=heading_path,
                    )
                )

            if options.include_tables:
                for table_i, table in enumerate(document.tables):
                    rows: list[list[str]] = []
                    for row in table.rows:
                        rows.append([cell.text.strip() for cell in row.cells])
                    if not any(cell for row in rows for cell in row):
                        continue
                    index += 1
                    blocks.append(
                        make_content_block(
                            seq_id=f"word-{index:04d}",
                            block_type=BlockType.TABLE,
                            text=table_to_text(rows),
                            table=rows,
                            location=SourceLocation(table_index=table_i),
                            heading_path=stack.path(),
                        )
                    )
        except Exception as exc:
            raise ArtifactParseError(f"Failed to extract Word file {path.name}: {exc}") from exc

        if not blocks:
            warnings.append("Word document contained no extractable paragraphs or tables.")

        return ArtifactDocument(
            source=str(path),
            artifact_type=ArtifactType.WORD,
            artifact_id=file_artifact_id(path),
            metadata=ArtifactMetadata(
                filename=path.name,
                artifact_type=ArtifactType.WORD,
                mime_type=MIME_BY_TYPE[ArtifactType.WORD],
                title=_meta_str(core.title),
                author=_meta_str(core.author),
                created=_meta_str(core.created),
                modified=_meta_str(core.modified),
            ),
            blocks=blocks,
            warnings=warnings,
        )


def _meta_str(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None
